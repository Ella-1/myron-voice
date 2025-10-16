from fastapi import FastAPI, UploadFile, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pathlib import Path
import os
import time
import shutil
from uuid import uuid4 

# REAL-WORLD IMPORTS for Transcription
import whisper
from docx import Document
from dotenv import load_dotenv
import logging

# Set up logging for better error diagnosis
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables (used for model configuration)
load_dotenv() 

# --- Configuration ---
# Use a temporary directory for all uploads to prevent file accumulation 
# on the main disk and ensure easier cleanup.
# NOTE: In a cloud environment, you would typically use S3/GCP Cloud Storage instead of local disk.
TEMP_ROOT_DIR = Path("/tmp/fastapi_uploads") 
TEMP_ROOT_DIR.mkdir(exist_ok=True)

# Configuration from environment variables
WHISPER_MODEL_NAME = os.getenv('WHISPER_MODEL', 'base')
# Get the source language hint (e.g., 'sw' for Swahili). Defaults to None for auto-detect.
SOURCE_LANGUAGE = os.getenv('SOURCE_LANGUAGE', None) 

# 1. Initialize FastAPI Application
app = FastAPI(title="FastAPI Transcriber Service (Multi-User Ready)")

# Setup CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# --- Model Loading and Utility Functions ---

# GLOBAL CONTAINER FOR LAZY LOADING: Ensures the large Whisper model is loaded only once
_MODEL_CONTAINER = [] 

# Define a persistent directory for the Whisper model download
# This ensures the model is downloaded only ONCE, not every time the server restarts.
PERSISTENT_MODEL_ROOT = Path("./.whisper_models")
PERSISTENT_MODEL_ROOT.mkdir(exist_ok=True)


def _load_whisper_model(model_name: str):
    """Loads the Whisper model if it hasn't been loaded yet in this process."""
    if not _MODEL_CONTAINER:
        try:
            logger.info(f"Worker process loading Whisper model: {model_name}...")
            # Use device="cpu" for server environments where GPUs might not be available
            # We now use the persistent directory defined above
            model = whisper.load_model(model_name, device="cpu", download_root=str(PERSISTENT_MODEL_ROOT.resolve()))
            _MODEL_CONTAINER.append(model)
            logger.info("Model loaded successfully.")
        except Exception as e:
            logger.error(f"FATAL ERROR: Could not load Whisper model '{model_name}': {e}")
            raise e
    return _MODEL_CONTAINER[0]

def secure_filename(filename: str) -> str:
    """A minimal secure filename implementation for safe path components."""
    return "".join(c for c in filename if c.isalnum() or c in ('.', '_', '-')).rstrip()

def cleanup_job_folder(job_dir: Path):
    """Utility function to run in a background task to safely delete the entire job directory."""
    if job_dir.exists():
        print(f"BACKGROUND TASK: Safely removing directory: {job_dir}")
        try:
            # Recursively remove the directory and all its contents
            shutil.rmtree(job_dir)
            print(f"BACKGROUND TASK: Cleanup complete for {job_dir}")
        except OSError as e:
            # Log cleanup errors but don't crash the server
            print(f"ERROR: Could not remove directory {job_dir}. Reason: {e}")

def transcribe_audio_and_save_doc(audio_filepath: Path, docx_filepath: Path):
    """
    Transcribes the audio file and saves the result as a DOCX file using Whisper.
    NOTE: FastAPI runs this blocking I/O/CPU-heavy function in a threadpool.
    """
    model = _load_whisper_model(WHISPER_MODEL_NAME) 
    
    # Prepare language parameters
    lang_param = None
    if SOURCE_LANGUAGE and SOURCE_LANGUAGE.strip():
        lang_param = SOURCE_LANGUAGE.strip().lower()

    try:
        # Resolve to string path for compatibility with the whisper library
        audio_filepath_str = str(audio_filepath.resolve())
        docx_filepath_str = str(docx_filepath.resolve())
        
        logger.info(f"Starting transcription/translation of {audio_filepath.name}...")
        
        # Transcribe and translate to English (task="translate")
        result = model.transcribe(
            audio_filepath_str,
            task="translate",
            language=lang_param,
            verbose=False 
        )
        
        transcript = result["text"]
        source_language_used = result.get('language', 'Unknown')
        
        # Create a Word document
        doc = Document()
        doc.add_heading("Audio Transcription (Translated to English)", level=1)
        doc.add_paragraph(f"Source File Processed: {audio_filepath.name}") 
        doc.add_paragraph(f"Source Language Detected/Hinted: {source_language_used.capitalize()}")
        
        if lang_param:
             doc.add_paragraph(f"Note: Explicit language hint '{lang_param}' was provided.")

        doc.add_paragraph("")
        doc.add_paragraph(transcript)

        # Save as Word file
        doc.save(docx_filepath_str) 
        logger.info(f"Transcription saved successfully to {docx_filepath_str}")

    except Exception as e:
        logger.error(f"Transcription failed for {audio_filepath.name}: {e}")
        # Re-raise the exception for the FastAPI endpoint to catch
        raise RuntimeError(f"Transcription Error: {e}") 


# --- The Multi-User Route ---

@app.post("/api/transcribe")
async def transcribe_and_download_endpoint(
    audio_file: UploadFile, 
    background_tasks: BackgroundTasks
):
    """Handles audio upload, transcription, and returns the resulting DOCX file."""

    if not audio_file.filename:
        raise HTTPException(status_code=400, detail="No file name provided.")

    # 1. Create a unique job ID and directory for this specific request/user
    job_id = uuid4()
    job_dir = TEMP_ROOT_DIR / str(job_id)
    job_dir.mkdir(parents=True, exist_ok=False)
    
    # Securely define filenames and paths within the unique directory
    safe_audio_filename = secure_filename(audio_file.filename)
    audio_filepath = job_dir / safe_audio_filename
    
    # Define the output path
    base_name = safe_audio_filename.split('.')[0]
    # Use a static filename to avoid conflicts during the actual transcription
    docx_filename = f"{base_name}_transcription.docx" 
    docx_filepath = job_dir / docx_filename

    # --- Start Processing ---
    try:
        # Save the uploaded file to the unique job directory
        contents = await audio_file.read()
        with open(audio_filepath, "wb") as f:
            f.write(contents)
        print(f"Saved input file to: {audio_filepath}")

        # Call the transcription function (FastAPI runs this blocking call in a thread pool)
        transcribe_audio_and_save_doc(audio_filepath, docx_filepath)
        
        # 2. Add the cleanup task to be executed AFTER the response is sent.
        background_tasks.add_task(cleanup_job_folder, job_dir)

        # 3. Send the file for download
        return FileResponse(
            path=docx_filepath,
            # Ensure the client sees the unique filename including the timestamp from the job_id
            filename=f"{base_name}_transcription_{int(time.time())}.docx",
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        
    except Exception as e:
        # 4. Handle errors and clean up immediately if the process failed *before* sending a response.
        logger.error(f"CRITICAL ERROR: Job {job_id} failed: {e}")
        
        # Immediate, blocking cleanup of the failed job directory before returning error response
        cleanup_job_folder(job_dir)
        
        return JSONResponse(
            status_code=500,
            content={'status': 'error', 'message': f'Transcription failed: {e}. Check server logs for job ID: {job_id}'}
        )
