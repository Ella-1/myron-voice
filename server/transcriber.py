import whisper
import os
from docx import Document
from dotenv import load_dotenv
import logging

# Set up logging for better error diagnosis in Gunicorn workers
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables (needed for model configuration)
load_dotenv() 

# --- Configuration ---
WHISPER_MODEL_NAME = os.getenv('WHISPER_MODEL', 'base')

# Get the source language hint (e.g., 'sw' for Swahili). Defaults to None for auto-detect.
SOURCE_LANGUAGE = os.getenv('SOURCE_LANGUAGE', None) 
# ---------------------

# GLOBAL CONTAINER FOR LAZY LOADING
# The model is heavy, so we load it only once per Gunicorn worker process,
# and not during the import phase, which prevents worker crashes.
_MODEL_CONTAINER = []

def _load_whisper_model(model_name):
    """Loads the Whisper model if it hasn't been loaded yet in this process."""
    if not _MODEL_CONTAINER:
        try:
            logger.info(f"Worker process loading Whisper model: {model_name}...")
            # Use fp16=False to reduce memory usage on CPU-only environments (common in standard web hosting)
            model = whisper.load_model(model_name, device="cpu", download_root="/tmp/whisper_models")
            _MODEL_CONTAINER.append(model)
            logger.info("Model loaded successfully.")
        except Exception as e:
            logger.error(f"FATAL ERROR: Could not load Whisper model '{model_name}': {e}")
            # Re-raise to ensure the worker fails with a clear message if loading is impossible
            raise e
    return _MODEL_CONTAINER[0]

def transcribe_audio_and_save_doc(audio_filepath: str, docx_filepath: str):
    """
    Transcribes the audio file and saves the result as a DOCX file.
    Designed to be called by the Flask endpoint.
    """
    model = _load_whisper_model(WHISPER_MODEL_NAME)
    
    # Prepare language parameters
    lang_param = None
    if SOURCE_LANGUAGE and SOURCE_LANGUAGE.strip():
        lang_param = SOURCE_LANGUAGE.strip().lower()

    try:
        logger.info(f"Starting transcription/translation of {audio_filepath}...")
        
        # Transcribe and translate to English (task="translate")
        result = model.transcribe(
            audio_filepath, 
            task="translate",
            language=lang_param,
            verbose=False # Suppress verbose output during request
        )
        
        transcript = result["text"]
        source_language_used = result.get('language', 'Unknown')
        
        # Create a Word document
        doc = Document()
        doc.add_heading("Audio Transcription (Translated to English)", level=1)
        doc.add_paragraph(f"Source File Processed: {os.path.basename(audio_filepath)}")
        doc.add_paragraph(f"Source Language Detected/Hinted: {source_language_used.capitalize()}")
        
        if lang_param:
             doc.add_paragraph(f"Note: Explicit language hint '{lang_param}' was provided.")

        doc.add_paragraph("")
        doc.add_paragraph(transcript)

        # Save as Word file
        doc.save(docx_filepath)
        logger.info(f"Transcription saved successfully to {docx_filepath}")

    except Exception as e:
        logger.error(f"Transcription failed for {audio_filepath}: {e}")
        # Re-raise the exception for the Flask endpoint to catch and return a 500 error
        raise Exception(f"Transcription Error: {e}")

# IMPORTANT: No code should run outside of functions here, except configuration and imports.
